import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING
from pathlib import Path
from typing import Dict, Any, List
import logging
import yaml

class WordConverter:
    def __init__(self, json_path: Path, output_path: Path, word_settings: Dict):
        self.json_path = json_path
        self.output_path = output_path
        self.word_settings = word_settings
        self.document = Document()
        self.circled_numbers = "①②③④⑤⑥⑦⑧⑨⑩"
        self.answers = []
        self.logger = self._setup_logger()

    def _setup_logger(self) -> logging.Logger:
        logger = logging.getLogger('word_converter')
        logger.setLevel(logging.INFO)
        
        # 로그 파일 경로 설정
        log_file = self.output_path.parent / 'word_converter.log'
        
        file_handler = logging.FileHandler(
            str(log_file),  # str로 변환하여 호출 오류 방지
            encoding='utf-8'
        )
        file_handler.setLevel(logging.INFO)
        
        console_handler = logging.StreamHandler()
        console_handler.setLevel(logging.INFO)
        
        formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
        file_handler.setFormatter(formatter)
        console_handler.setFormatter(formatter)
        
        logger.addHandler(file_handler)
        logger.addHandler(console_handler)
        
        return logger

    def load_json_data(self) -> List[Dict]:
        with open(self.json_path, 'r', encoding='utf-8') as file:
            return json.load(file)

    def set_font(self, run, font_settings: Dict):
        font = run.font
        font.name = font_settings['name']
        font.size = Pt(font_settings['size'])
        if 'character_spacing' in font_settings:
            font.spacing = Pt(font_settings['character_spacing'])

    def set_paragraph_format(self, paragraph, line_spacing: float):
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph_format.line_spacing = Pt(line_spacing)

    def add_paragraph_with_format(self, text: str, font_type: str):
        font_settings = self.word_settings['fonts'][font_type]
        paragraph = self.document.add_paragraph()
        self.set_paragraph_format(paragraph, font_settings['line_spacing'])
        run = paragraph.add_run(text)
        self.set_font(run, font_settings)

    def number_to_circled(self, number: int) -> str:
        return self.circled_numbers[int(number) - 1]

    def convert_to_word(self) -> bool:
        """JSON을 Word 문서로 변환"""
        try:
            data = self.load_json_data()
            self.logger.info(f"JSON 데이터 로드 완료: {len(data)}개 지문")

            for item in data:
                self.add_paragraph_with_format(f"{item['본문번호']}.", 'passage')
                self.add_paragraph_with_format(item['본문'], 'passage')
                self.document.add_paragraph()

                for problem in item['문제들']:

                    # 잘못된 str 호출 오류 수정
                    self.add_paragraph_with_format(
                        f"{problem['문제번호']}. {problem['문제']}", 
                        'question'
                    )
                    
                    for i, option in enumerate(problem['선지']):
                        # 인덱스에 맞춰 옵션을 추가
                        self.add_paragraph_with_format(
                            f"{self.number_to_circled(i + 1)} {option[3:]}", 
                            'option'
                        )
                    
                    self.document.add_paragraph()
                    answer_index = len(self.answers) % len(item['정답'])
                    circled_answer = self.number_to_circled(int(item['정답'][answer_index]))  # 정답도 int로 변환 후 처리
                    self.answers.append(f"{problem['문제번호']}.{circled_answer}")

            # 정답 추가
            self.document.add_page_break()
            for answer in self.answers:
                self.add_paragraph_with_format(answer, 'answer')

            # 페이지 여백 설정
            margins = self.word_settings['margins']
            sections = self.document.sections
            for section in sections:
                section.top_margin = Cm(margins['top'])
                section.bottom_margin = Cm(margins['bottom'])
                section.left_margin = Cm(margins['left'])
                section.right_margin = Cm(margins['right'])

            # 문서 저장
            self.output_path.parent.mkdir(parents=True, exist_ok=True)
            self.document.save(str(self.output_path))
            
            self.logger.info(f"Word 문서 생성 완료: {self.output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Word 문서 생성 중 오류 발생: {str(e)}")
            return False
        
    @classmethod
    def from_config(cls, config_path: str = "config.yaml"):
        """설정 파일에서 WordConverter 인스턴스 생성"""
        with open(config_path, 'r', encoding='utf-8') as f:
            config = yaml.safe_load(f)

        # 기본 경로 설정
        base_dir = Path(config_path).parent
        generated_problems_dir = base_dir / "output/generated_problems"
        word_output_dir = base_dir / "output/word_documents"

        # JSON 파일 찾기 (가장 최근 파일)
        json_files = list(generated_problems_dir.glob("*.json"))
        if not json_files:
            raise FileNotFoundError(f"JSON 파일을 찾을 수 없습니다: {generated_problems_dir}")
        
        json_path = max(json_files, key=lambda x: x.stat().st_mtime)
        word_output = word_output_dir / f"{json_path.stem}.docx"

        # 출력 디렉토리 생성
        word_output_dir.mkdir(parents=True, exist_ok=True)

        return cls(
            json_path=json_path,
            output_path=word_output,
            word_settings=config.get('word_settings', {})
        )

def main():
    try:
        print("Word 문서 변환 시작...")
        
        # 현재 디렉토리에서 config.yaml 찾기
        config_path = Path(__file__).parent / "config.yaml"
        if not config_path.exists():
            print(f"Error: 설정 파일을 찾을 수 없습니다: {config_path}")
            return

        # WordConverter 인스턴스 생성
        converter = WordConverter.from_config(str(config_path))
        
        # 변환 실행
        if converter.convert_to_word():
            print("Word 문서 변환이 완료되었습니다.")
            print(f"출력 파일: {converter.output_path}")
        else:
            print("Word 문서 변환 중 오류가 발생했습니다.")
        
    except Exception as e:
        print(f"Error: {str(e)}")

if __name__ == "__main__":
    main()
